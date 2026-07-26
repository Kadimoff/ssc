#!/usr/bin/env python3
"""Generate the branded SSC partnership strategy DOCX from the preserved research Markdown."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/research/deep-research-report.md"
OUTPUT = ROOT / "docs/SSC_QS_Power_of_Partnerships_Transformation_Report.docx"

NAVY = "111827"
NAVY_2 = "172033"
EMERALD = "10B981"
EMERALD_DARK = "047857"
GOLD = "F5B840"
MUTED = "64748B"
PALE = "F1F5F9"
WHITE = "FFFFFF"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color: str = "CBD5E1", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_hyperlink(paragraph, text: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), EMERALD_DARK)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    run.append(props)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(EMERALD_DARK)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def clean_research_text(text: str) -> str:
    text = re.sub(r"\s*(?:filecite|cite)[^]+", "", text)
    return text.replace("\u00a0", " ").rstrip()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.12

    for level, size, color in ((1, 24, NAVY), (2, 18, EMERALD_DARK), (3, 14, NAVY), (4, 11.5, EMERALD_DARK)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if level > 1 else 20)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in document.styles:
        code = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(8)
        code.font.color.rgb = RGBColor.from_string(WHITE)
        code.paragraph_format.left_indent = Inches(0.14)
        code.paragraph_format.right_indent = Inches(0.14)
        code.paragraph_format.space_before = Pt(3)
        code.paragraph_format.space_after = Pt(3)

    if "Callout" not in document.styles:
        callout = document.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Aptos"
        callout.font.size = Pt(10)
        callout.font.italic = True
        callout.font.color.rgb = RGBColor.from_string(NAVY)
        callout.paragraph_format.left_indent = Inches(0.2)
        callout.paragraph_format.right_indent = Inches(0.1)
        callout.paragraph_format.space_before = Pt(6)
        callout.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SSC  •  PARTNERSHIP OPERATING SYSTEM")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(EMERALD_DARK)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SSC Transformation Report   •   ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(footer, "PAGE", "1")

    core = document.core_properties
    core.title = "SSC QS Power of Partnerships Transformation Report"
    core.subject = "Partnership-first product, architecture, pilot, and QS strategy"
    core.author = "Student Startup Community"
    core.keywords = "SSC, partnerships, student entrepreneurship, QS Reimagine Education"


def add_cover(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("STUDENT STARTUP COMMUNITY")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(EMERALD_DARK)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("QS Power of Partnerships\nTransformation Report")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━━━━━━━━━━━━━━━━━━━━")
    run.font.color.rgb = RGBColor.from_string(GOLD)
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "From verified founder community to auditable partnership operating system"
    )
    run.font.name = "Aptos"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    for _ in range(4):
        document.add_paragraph()

    meta = document.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    labels = (("Direction", "Partnership-first, founder workflows preserved"), ("Delivery", "Five implementation parts"), ("Status", "Implementation strategy • 23 July 2026"))
    for row, (label, value) in zip(meta.rows, labels):
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(4.85)
        set_cell_fill(row.cells[0], NAVY)
        set_cell_fill(row.cells[1], PALE)
        set_cell_border(row.cells[0], NAVY)
        set_cell_border(row.cells[1], "E2E8F0")
        left = row.cells[0].paragraphs[0]
        left_run = left.add_run(label.upper())
        left_run.bold = True
        left_run.font.size = Pt(8)
        left_run.font.color.rgb = RGBColor.from_string(WHITE)
        right = row.cells[1].paragraphs[0]
        right_run = right.add_run(value)
        right_run.font.size = Pt(9.5)
        right_run.font.color.rgb = RGBColor.from_string(NAVY)

    document.add_paragraph()
    notice = document.add_paragraph(style="Callout")
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.add_run(
        "Internal strategy document. Prototype and planning figures are illustrative until verified by real partner-backed pilot evidence."
    )
    shade_paragraph(notice, "ECFDF5", EMERALD)
    document.add_page_break()

    heading = document.add_heading("Contents", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc = document.add_paragraph()
    add_field(toc, r'TOC \o "1-3" \h \z \u', "Right-click and update field to refresh the table of contents.")
    document.add_page_break()


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(normalized):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_fill(cell, NAVY if row_index == 0 else (PALE if row_index % 2 == 0 else WHITE))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value.strip())
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    previous: list[str] | None = None
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        values = [clean_research_text(cell.strip()) for cell in raw.split("|")]
        is_separator = all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in values)
        if not is_separator and values != previous:
            rows.append(values)
            previous = values
        index += 1
    return rows, index


def add_source_notes(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Source provenance and verification notes", level=1)
    paragraph = document.add_paragraph()
    add_inline(
        paragraph,
        "The preserved source is `docs/research/deep-research-report.md`. Tool-specific citation markers were removed from this polished edition because they are not portable references.",
    )
    document.add_paragraph(
        "Decision-critical QS dates and application statements must be checked against the live official page before submission."
    )
    links = (
        ("QS Reimagine Education 2026 — applications", "https://www.qs.com/conferences/reimagine/apply"),
        ("QS Reimagine Education 2026 — conference", "https://www.qs.com/conferences/reimagine"),
        ("PostgreSQL JSON types", "https://www.postgresql.org/docs/current/datatype-json.html"),
        ("OWASP Application Security Verification Standard", "https://owasp.org/www-project-application-security-verification-standard/"),
        ("WCAG 2.2", "https://www.w3.org/TR/WCAG22/"),
        ("WebAuthn", "https://www.w3.org/TR/webauthn-3/"),
    )
    for label, url in links:
        paragraph = document.add_paragraph(style="List Bullet")
        add_hyperlink(paragraph, label, url)


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    skipped_title = False

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_language = stripped[3:].strip()
                code_lines = []
            else:
                label = "Diagram specification" if code_language == "mermaid" else (code_language.upper() if code_language else "CODE")
                heading = document.add_paragraph()
                run = heading.add_run(label)
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string(EMERALD_DARK)
                paragraph = document.add_paragraph(style="Code Block")
                paragraph.add_run("\n".join(code_lines))
                shade_paragraph(paragraph, NAVY_2, EMERALD if code_language == "mermaid" else None)
                in_code = False
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = clean_research_text(heading_match.group(2))
            if level == 1 and not skipped_title:
                skipped_title = True
            else:
                document.add_heading(title, level=min(level, 4))
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        text = clean_research_text(stripped)
        if text.startswith("> "):
            paragraph = document.add_paragraph(style="Callout")
            add_inline(paragraph, text[2:])
            shade_paragraph(paragraph, "ECFDF5", EMERALD)
        elif re.match(r"^[-*]\s+", text):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, re.sub(r"^[-*]\s+", "", text))
        elif re.match(r"^\d+\.\s+", text):
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", text))
        else:
            paragraph = document.add_paragraph()
            add_inline(paragraph, text)
        index += 1


def main() -> int:
    if not SOURCE.exists():
        print(f"Missing source: {SOURCE}", file=sys.stderr)
        return 1
    document = Document()
    configure_document(document)
    add_cover(document)
    render_markdown(document, SOURCE.read_text(encoding="utf-8"))
    add_source_notes(document)
    document.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
